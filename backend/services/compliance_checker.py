"""
Procurement compliance checker.

Tender side:
  1. Extract text (pdfplumber / python-docx)
  2. Find requirements section via keyword list
  3. Segment into text units
  4. RoBERTa binary classifier — label 1 = requirement (threshold 0.5)

Supplier side:
  1. Extract text from supplier document
  2. Encode requirements + supplier sentences with all-MiniLM-L6-v2
  3. Cosine similarity → Met ≥0.50 / Partial 0.35–0.50 / Not Met <0.35
  4. Negation detection — downgrade to Not Met
  5. Mandatory check — disqualify if Not Met on license/registration/tax
  6. Score = (Met + Partial×0.5) / Total × 100%
"""

import re
from pathlib import Path

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


REQUIREMENTS_SECTION_KEYWORDS = [
    'requirements', 'eligibility', 'criteria', 'mandatory', 'qualification',
    'experience required', 'documentation', 'documents required',
    'minimum requirements', 'technical requirements', 'financial requirements',
    'conditions', 'must', 'shall',
]

MANDATORY_KEYWORDS = [
    'license', 'licence', 'registration', 'tax compliance', 'tax clearance',
    'business registration', 'certified', 'certification', 'permit',
    'authority', 'mandatory', 'compulsory', 'required by law',
]

NEGATION_PHRASES = [
    'no experience', 'not licensed', 'not certified', 'not registered',
    'no license', 'no licence', 'no registration', 'not qualified',
    'no certification', 'lack of', 'unable to', 'do not have',
    'does not have', 'without certification', 'without license',
]


# ─────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────

def extract_text(file_path):
    ext = Path(file_path).suffix.lower()
    if ext == '.pdf':
        if not PDF_SUPPORT:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        with pdfplumber.open(file_path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages).strip()
    elif ext in ('.docx', '.doc'):
        if not DOCX_SUPPORT:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        doc = DocxDocument(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─────────────────────────────────────────────────────────────
# TEXT SEGMENTATION
# ---

def _find_requirements_section(text):
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if any(kw in line.lower() for kw in REQUIREMENTS_SECTION_KEYWORDS):
            return '\n'.join(lines[i:])
    return text


_NEW_ITEM = re.compile(r'^(\d+[\.)\]]\s+|[a-zA-Z][\.)\]]\s+|[-*]\s+)')

def _ends_sentence(line):
    return bool(re.search(r'[.!?;:]\s*$', line))


def _segment(text):
    raw_lines = text.split('\n')
    cleaned = []
    for line in raw_lines:
        stripped = line.strip()
        if not stripped or stripped.isdigit() or stripped.startswith('--- PAGE'):
            continue
        cleaned.append(stripped)

    units = []
    current = ''
    for line in cleaned:
        if not current:
            current = line
            continue

        # Rule 1: explicit continuation
        if line[0].islower() or line.startswith(('and ', 'or ', 'including ', 'such as ', 'as well as ', 'with ')):
            current += ' ' + line
            continue

        # Rule 2: previous line ended a sentence
        if _ends_sentence(current):
            units.append(current)
            current = line
            continue

        # Rule 3: new list item
        if _NEW_ITEM.match(line):
            units.append(current)
            current = line
            continue

        # Rule 4: unit already long
        if len(current.split()) > 40:
            units.append(current)
            current = line
            continue

        # Rule 5: mid-sentence wrap
        current += ' ' + line

    if current:
        units.append(current)

    result = []
    for block in units:
        if len(block.split()) <= 60:
            result.append(block)
        else:
            sents = re.split(r'(?<=[.!?])\s+(?=[A-Z])', block)
            result.extend(sents)

    return [u.strip() for u in result if u.strip() and len(u.split()) >= 3]

# TENDER SIDE — RoBERTa requirement extraction
# ─────────────────────────────────────────────────────────────

def extract_requirements(roberta_model, tokenizer, text, threshold=0.87):
    """
    Classify each text unit as requirement (label 1) or not (label 0).
    Returns list of dicts: {text, confidence, is_mandatory}

    Uses two-pass strategy: first strict threshold (0.87), then relaxed (0.55)
    if the first pass yields nothing, to handle varied document formats.
    """
    import torch

    section = _find_requirements_section(text)
    units = _segment(section)
    if not units:
        return []

    scored = []
    for unit in units:
        inputs = tokenizer(
            unit,
            max_length=128,
            padding='max_length',
            truncation=True,
            return_tensors='pt',
        )
        with torch.no_grad():
            outputs = roberta_model(**inputs)
            probs = torch.softmax(outputs.logits, dim=1).cpu().numpy()[0]
        scored.append((unit, float(probs[1])))

    # First pass: strict threshold
    results = [
        {
            'text':         unit,
            'confidence':   round(prob, 4),
            'is_mandatory': any(kw in unit.lower() for kw in MANDATORY_KEYWORDS),
        }
        for unit, prob in scored if prob >= threshold
    ]

    # Second pass: relaxed threshold — only used when nothing passed the strict gate
    if not results:
        FALLBACK = 0.55
        results = [
            {
                'text':         unit,
                'confidence':   round(prob, 4),
                'is_mandatory': any(kw in unit.lower() for kw in MANDATORY_KEYWORDS),
            }
            for unit, prob in scored if prob >= FALLBACK
        ]

    return results


# ─────────────────────────────────────────────────────────────
# SUPPLIER SIDE — sentence-transformers matching
# ─────────────────────────────────────────────────────────────

def match_supplier(sentence_model, requirements, supplier_text):
    """
    Match supplier document against a list of requirements.

    Args:
        sentence_model   SentenceTransformer instance
        requirements     list of dicts {text, is_mandatory, ...} OR plain strings
        supplier_text    full text of supplier proposal

    Returns:
        items           list of per-requirement result dicts
        total_score     float 0–100
        disqualified    bool
        disq_reasons    list of str
    """
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np

    if not requirements:
        return [], 0.0, False, []

    def _all_missing(reqs):
        """Return every requirement as Not Met when the proposal has no extractable text."""
        items = []
        disq = False
        disq_r = []
        for r in reqs:
            req_text = r['text'] if isinstance(r, dict) else r
            is_mand  = r.get('is_mandatory', False) if isinstance(r, dict) else False
            if is_mand:
                disq = True
                disq_r.append(req_text[:100])
            items.append({'requirement_text': req_text, 'status': 'Not Met',
                          'similarity': 0.0, 'best_match': '', 'is_mandatory': is_mand})
        return items, 0.0, disq, disq_r

    if not supplier_text:
        return _all_missing(requirements)

    supplier_sentences = _segment(supplier_text)
    if not supplier_sentences:
        return _all_missing(requirements)

    req_texts = [r['text'] if isinstance(r, dict) else r for r in requirements]
    req_embeddings = sentence_model.encode(req_texts)
    sup_embeddings = sentence_model.encode(supplier_sentences)

    items = []
    total_weight = len(requirements)
    total_earned = 0.0
    disqualified = False
    disq_reasons = []

    for i, req in enumerate(requirements):
        req_text = req['text'] if isinstance(req, dict) else req
        is_mandatory = req.get('is_mandatory', False) if isinstance(req, dict) else False

        sims = cosine_similarity([req_embeddings[i]], sup_embeddings)[0]
        best_idx = int(np.argmax(sims))
        best_score = float(sims[best_idx])
        best_sentence = supplier_sentences[best_idx]

        # Negation detection
        if any(phrase in best_sentence.lower() for phrase in NEGATION_PHRASES):
            best_score = 0.0

        if best_score >= 0.50:
            status = 'Met'
            total_earned += 1.0
        elif best_score >= 0.35:
            status = 'Partial'
            total_earned += 0.5
        else:
            status = 'Not Met'
            if is_mandatory:
                disqualified = True
                disq_reasons.append(req_text[:100])

        items.append({
            'requirement_text': req_text,
            'status':           status,
            'similarity':       round(best_score, 4),
            'best_match':       best_sentence[:200],
            'is_mandatory':     is_mandatory,
        })

    score = round(total_earned / total_weight * 100, 1) if total_weight > 0 else 0.0
    return items, score, disqualified, disq_reasons
